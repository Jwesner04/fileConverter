#!/usr/bin/env python
# -*- coding: cp1252 -*-
# -*- coding: utf-8 -*-
# ## asciiToDoc.py ############################################################
#
# Code Description: Inputs files and outputs a summary to a docx file.
# Last Edited: 10/25/2016
# Last Edited By: Jonathan Wesner
# Last Changes Made: ...
#
# #############################################################################
import sys, re, os, urllib2, string, datetime, docx, lxml.etree, jinja2
from docx.shared import Pt, Inches
from docxtpl import DocxTemplate
from PyQt4 import QtCore, QtGui, uic
import images_rc

# Include user interface files
cwd = os.getcwd()
qtMainWindowFile = cwd + "/interface.ui"
qtMyPopUpFile = cwd + "/myPopUp.ui"

Ui_MainWindow, QtBaseClass = uic.loadUiType(qtMainWindowFile)
Ui_PopUpWindow, QtSubClass = uic.loadUiType(qtMyPopUpFile)


# This is a template class for any popup errors that need to be made
# ----------------------------------------------------------------------------#
#
# Class Description: This class is used for error checking and messages
#                    to user.
# Last Edited: 10/13/2016
# Last Edited By: Jonathan Wesner
# Last Changes Made: ...
#
# ----------------------------------------------------------------------------#

class MyPopup(QtGui.QWidget, Ui_PopUpWindow):
    def __init__(self):
        QtGui.QWidget.__init__(self)
        Ui_PopUpWindow.__init__(self)
        self.setupUi(self)

        # ---------------------------------------------------#
        # User Interface Property changes
        # ---------------------------------------------------#
        self.errorButton.setCursor(QtGui.QCursor(QtCore.Qt.PointingHandCursor))

        # ---------------------------------------------------#
        # Call appropriate function on user actions
        # ---------------------------------------------------#
        # Connects to startSystem() function
        self.errorButton.clicked.connect(self.close)


# ----------------------------------------------------------------------------#
#
# Class Description: This class is used for the main window.
# Last Edited: 10/13/2016
# Last Edited By: Jonathan Wesner
# Last Changes Made: ...
#
# ----------------------------------------------------------------------------#

class MainWindow(QtGui.QDialog, Ui_MainWindow):
    def __init__(self):
        QtGui.QMainWindow.__init__(self)
        Ui_MainWindow.__init__(self)
        self.setupUi(self)
        self.errorPopup = MyPopup()

        # ---------------------------------------------------#
        # Local Variables
        # ---------------------------------------------------#
        # Create instance of Document.docx file
        self.filePath = ""
        self.userName = ""
        self.dateNow = datetime.datetime.now().date().strftime("%m-%d-%Y")
        self.projectName = ""
        self.document = DocxTemplate(cwd + "/default.docx")
        self.tableRowCount = 0

        # ---------------------------------------------------#
        # file name Variables -- add to list if additional
        #                       files are parsed in future
        # ---------------------------------------------------#
        self.fsum = ""
        self.fsum_erc = ""
        self.fcls = ""
        self.flog = ""
        self.fascii = ""
        self.fpictures = []

        # ---------------------------------------------------#
        # User Interface Property changes
        # ---------------------------------------------------#
        self.date.setText("Date: " + self.dateNow)

        # ---------------------------------------------------#
        # Call appropriate function on user actions
        # ---------------------------------------------------#
        self.okButton.clicked.connect(self.systemHandler)
        self.toolButton.clicked.connect(self.openExplorer)
        self.cancelButton.clicked.connect(self.closeApplication)

    # ----------------------------------------------------------------------------#
    #
    # Function Name: closeApplication()
    # Description: Called whenever cancel button is clicked by user. The
    #              application will close
    # Last Edited: 11/05/2016
    # Last Edited By: Jonathan Wesner
    # Last Changes Made: ...
    #
    # ----------------------------------------------------------------------------#

    def closeApplication(self):
        sys.exit(app.exec_())

    def openExplorer(self):
        fileDirectory = QtGui.QFileDialog.getExistingDirectory(self, "Select Directory")
        self.fileInput.setText(fileDirectory)

    def systemHandler(self):
        # ------------------------------------------------------------------------#
        # Handle file path error checking
        # ------------------------------------------------------------------------#
        self.filePath = str(self.fileInput.text())
        # Check that file path is valid
        if not os.path.exists(os.path.dirname(self.filePath)):
            self.displayError("File path provided not found.")
            return

        # ------------------------------------------------------------------------#
        # Handle user Name and Project Name and check that user has not left
        # blank.
        # ------------------------------------------------------------------------#
        self.userName = str(self.userNameEdit.text())
        self.projectName = str(self.projectNameEdit.text())
        if (self.userName is "") or (self.projectName is ""):
            self.displayError("Please fill all fields including User Name and Project Name.")
            return

        # ------------------------------------------------------------------------#
        # Set title information and styling of Word Document here
        # ------------------------------------------------------------------------#
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)
        self.printToDocument('Project Name: ' + self.projectName + '\n' + 'Tape-out Date: ' +
                             str(self.dateNow) + '\n' + 'Preparer: ' + self.userName + '\n\n', True)

        # ------------------------------------------------------------------------#
        # Add Title Table
        #   table.cell(Row,Column) --- reference
        # ------------------------------------------------------------------------#
        self.tableRowCount = 0
        self.table = self.document.add_table(rows=0, cols=4)
        self.table.style = 'Light Shading'
        self.fillTable('File name', 'Description', 'File size (bytes)', 'Date/Time')
        # parse files and fill table
        self.getFileNames()
        self.document.add_page_break()

        # ------------------------------------------------------------------------#
        # Error check all files are present before continuing
        # ------------------------------------------------------------------------#
        if self.fsum == "":
            self.displayError("No .sum file exists. Please include file in folder before continuing.")
            return
        if self.fsum_erc == "":
            self.displayError("No .sum_erc file exists. Please include file in folder before continuing.")
            return
        if self.fcls == "":
            self.displayError("No .rep.cls file exists. Please include file in folder before continuing.")
            return
        if self.fascii == "":
            self.displayError("No .drc_errors.ascii file exists. Please include file in folder before continuing.")
            return
        if self.flog == "":
            self.displayError("No .sum_erc file exists. Please include file in folder before continuing.")
            return

        # ------------------------------------------------------------------------#
        #   Order of Looking for files -- add to list as needed
        #       1)  .sum file
        #       2)  .ascii file (used to output errors of .sum file)
        #       3)  .cls file
        #       4)  .sum_erc
        #       5)  .log
        #       6)  .png or jpg (add pictures last)
        # ------------------------------------------------------------------------#
        collectErrors = ""
        # first argument of searchSumFile() takes in file name, second argument
        # allows for the collection of the summary section of the file
        self.printToDocument('DRC:', True)
        # Order File look up #1
        collectErrors = self.searchSumFile(self.fsum, False)

        # Order File look up #2
        self.printToDocument("DRC Errors:", True)
        self.sumFileErrors(collectErrors, self.fascii)

        # Order File look up #3
        self.printToDocument("LVS:", True)
        self.searchClsFile(self.fcls)

        # Order File look up #4
        collectErrors = ""
        self.printToDocument("ERC:", True)
        collectErrors = self.searchSumFile(self.fsum_erc, True)

        # Order File look up #5
        self.printToDocument("ERC Errors:", True)
        self.sumErcFileErrors(collectErrors, self.flog)

        # Add pictures here...order file look up #6
        self.printToDocument("Images:", True)
        if self.fpictures is not None:
            for img in self.fpictures:
                self.document.add_picture(img, width=Inches(3), height=Inches(3))

        # Save document
        self.document.save(self.filePath + '/summary.docx')

        # Message to user that file is done
        self.displayError("File is Complete! File saved in folder where other files were specified.", "ATTENTION!")
        self.fileInput.setText("")
        self.userNameEdit.setText("")
        self.projectNameEdit.setText("")

    # ----------------------------------------------------------------------------#
    #
    # Function Name: getFileNames()
    # Description: This function takes care of grabing the file names needed to be
    #              opened. It uses the filePath variable to look for the files.
    #              Current files being read in are in this order...
    #                   1)  .sum file
    #                   2)  .ascii file (used to output errors of .sum file)
    #                   3)  .cls file
    #                   4)  .sum_erc file
    #                   5)  .log file
    #                   6)  .png or jpg(add pictures last)
    # Last Edited: 11/05/2016
    # Last Edited By: Jonathan Wesner
    # Last Changes Made: ...
    #
    # ----------------------------------------------------------------------------#

    def getFileNames(self):
        files = os.walk(self.filePath, topdown=False)
        for fileName in os.listdir(self.filePath):
            # --------------------------------------------------------------------#
            # get all files in folder and print out info to a table
            # --------------------------------------------------------------------#
            f = os.path.join(self.filePath, fileName)
            fileStats = os.stat(f)
            fileSize = fileStats.st_size
            dateLastModified = fileStats.st_mtime
            timeFormatted = datetime.datetime.fromtimestamp(int(dateLastModified)).strftime('%Y-%m-%d %H:%M:%S')
            self.fillTable(fileName, "", fileSize, timeFormatted)
            # --------------------------------------------------------------------#
            # check what type fileName is and fill the variables
            # --------------------------------------------------------------------#
            if fileName.endswith(".sum_erc"):
                self.fsum_erc = os.path.join(self.filePath, fileName)
            elif fileName.endswith(".sum"):
                self.fsum = os.path.join(self.filePath, fileName)
            elif fileName.endswith(".log"):
                if fileName.endswith(".streamout"):
                    return
                self.flog = os.path.join(self.filePath, fileName)
            elif fileName.endswith(".rep.cls"):
                self.fcls = os.path.join(self.filePath, fileName)
            elif fileName.endswith(".drc_errors.ascii"):
                self.fascii = os.path.join(self.filePath, fileName)
            else:
                if fileName.endswith(".png") or fileName.endswith(".jpg"):
                    self.fpictures.append(os.path.join(self.filePath, fileName))

    # ----------------------------------------------------------------------------#
    #
    # Function Name: fillTable()
    # Description: This function takes in four data strings from each file
    #              that will be parsed and fills a table with the individual
    #              attributes of each file...
    # Last Edited: 10/29/2016
    # Last Edited By: Jonathan Wesner
    # Last Changes Made: ...
    #
    # ----------------------------------------------------------------------------#

    def fillTable(self, fileName, fileDescription, fileSize, fileDateTime):
        # Add Row
        self.table.add_row()

        # Define each column of the table before setting the values
        fName = self.table.cell(self.tableRowCount, 0)
        fDescription = self.table.cell(self.tableRowCount, 1)
        fSize = self.table.cell(self.tableRowCount, 2)
        fDateTime = self.table.cell(self.tableRowCount, 3)

        # fill each cell from passed in strings
        fName.text = fileName
        fDescription.text = fileDescription
        fSize.text = str(fileSize)
        fDateTime.text = fileDateTime

        # Add to tableRowCount
        self.tableRowCount = self.tableRowCount + 1

    # ----------------------------------------------------------------------------#
    #
    # Function Name: printToDocument()
    # Description: This function handles creating a new paragraph to the document.
    #              It takes in a string and will pass that string to the document.
    # Last Edited: 10/27/2016
    # Last Edited By: Jonathan Wesner
    # Last Changes Made: ...
    #
    # ----------------------------------------------------------------------------#

    def printToDocument(self, string, isBold=False):
        output = self.document.add_paragraph()
        output.add_run(string).bold = isBold

    # ----------------------------------------------------------------------------#
    #
    # Function Name: searchSumFile()
    # Description: This function handles searching the .sum file. It parses
    #              through the header of the file, and finds any errors under the
    #              RULECHECK section of the file. These errors are collected and
    #              passed to the sumFileErrors() function, which outputs the
    #              details of the errors from the drc_errors.ascii file.
    # Last Edited: 10/27/2016
    # Last Edited By: Jonathan Wesner
    # Last Changes Made: ...
    #
    # ----------------------------------------------------------------------------#

    def searchSumFile(self, fileToSearch, getSummary):

        sumFileHeader = ""
        sumFileSummary = ""
        sumHeaderDone = False
        summarySection = False
        sumCollectErrors = ""
        headerCount = 0

        # ------------------------------------------------------------------------#
        # Open .sum file and parse through to grab header/summary, and any errors
        # ------------------------------------------------------------------------#
        sumFile = open(fileToSearch, 'r')
        for line in sumFile:
            if (('------' in line) or ('******' in line)) and (sumHeaderDone is False):
                headerCount = headerCount + 1
                if headerCount > 1:
                    sumHeaderDone = True
                    sumFileHeader = sumFileHeader + line
            # Check for header finish
            if sumHeaderDone is False:
                sumFileHeader = sumFileHeader + line
            # Check for errors in file
            if sumHeaderDone:
                if ("RULECHECK " in line) and ("Total Result      0 (       0)" not in line):
                    sumCollectErrors = sumCollectErrors + " " + re.search('RULECHECK (.+?) ....', line).group(1) + " "
            # Parse through summary of file
            if getSummary:
                if "--- SUMMARY" in line:
                    summarySection = True
                if summarySection:
                    sumFileSummary = sumFileSummary + line

        # print header to document
        self.printToDocument(sumFileHeader)

        # print summary to header if needed
        if getSummary:
            self.printToDocument(sumFileSummary)

        # Close file
        sumFile.close()
        return sumCollectErrors

    # ----------------------------------------------------------------------------#
    #
    # Function Name: sumFileErrors()
    # Description: This function is called by searchSumFile() and outputs
    #              details about the particular errors found in the
    #              searchSumFile().
    # Last Edited: 10/27/2016
    # Last Edited By: Jonathan Wesner
    # Last Changes Made: ...
    #
    # ----------------------------------------------------------------------------#

    def sumFileErrors(self, errorsToSearch, errorFileToSearch):

        checkStatus = False
        sumErrors = ""
        errorFile = open(errorFileToSearch, 'r')
        for line in errorFile:

            # remove '\n' at end of line
            stripLine = line.rstrip()
            result = re.search(stripLine, errorsToSearch)
            if checkStatus is True:
                sumErrors = sumErrors + line
                if error in stripLine:
                    sumErrors = sumErrors + "\n"
                    checkStatus = False

            if result is not None:
                checkStatus = True
                error = stripLine
                sumErrors = sumErrors + line

        self.printToDocument(sumErrors)
        errorFile.close()

    # ----------------------------------------------------------------------------#
    #
    # Function Name: sumFileErrors()
    # Description: This function is called by searchSumFile() and outputs
    #              details about the particular errors found in the
    #              .sum_erc file
    # Last Edited: 10/29/2016
    # Last Edited By: Jonathan Wesner
    # Last Changes Made: ...
    #
    # ----------------------------------------------------------------------------#

    def sumErcFileErrors(self, errorsToSearch, errorFileToSearch):

        sumErrors = ""
        errorFile = open(errorFileToSearch, 'r')
        errorFound = False
        errorsToSearch = re.findall('ERC(.+?) ', errorsToSearch)
        for line in errorFile:

            # remove '\n' at end of line
            if errorFound:
                sumErrors = sumErrors + line
                if "}" in line:
                    errorFound = False
            else:
                for error in errorsToSearch:
                    if (('RULE ERC' + error + " {") in line):
                        sumErrors = sumErrors + line
                        errorFound = True
                        break

        self.printToDocument(sumErrors)
        errorFile.close()

    # ----------------------------------------------------------------------------#
    #
    # Function Name: searchClsFile()
    # Description: This function is searches for and prints out the header
    #              in the .cls file.
    # Last Edited: 10/27/2016
    # Last Edited By: Jonathan Wesner
    # Last Changes Made: ...
    #
    # ----------------------------------------------------------------------------#

    def searchClsFile(self, fileToSearch):

        clsHeaderDone = False
        clsFileHeader = ""
        headerCount = 0
        clsFile = open(fileToSearch, 'r')

        for line in clsFile:
            if ('###########' in line) and (clsHeaderDone == False):
                headerCount = headerCount + 1
                if headerCount > 1:
                    clsHeaderDone = True
                    clsFileHeader = clsFileHeader + line

            if clsHeaderDone == False:
                clsFileHeader = clsFileHeader + line

        self.printToDocument(clsFileHeader)
        clsFile.close()

    # #############################################################################
    #
    # Function Description: This function is called whenever an error message
    #                       needs to be displayed to the screen.
    # Last Edited: 11/12/2016
    # Last Edited By: Jonathan Wesner
    # Last Changes Made: ...
    #
    # #############################################################################

    def displayError(self, errorText, title="ERROR!"):
        self.errorPopup.errorMessage.setText(errorText)
        self.errorPopup.title.setText(title)
        # Freezes the Main Window till a response is made by the user for MyPopup()
        self.errorPopup.setWindowModality(QtCore.Qt.ApplicationModal)
        self.errorPopup.show()


# ----------------------------------------------------------------------------#
#
# Main Description: The main parses through files within folder and output
#                   summary of each file goes onto a docx file.
# Last Edited: 10/25/2016
# Last Edited By: Jonathan Wesner
# Last Changes Made: ...
#
# ----------------------------------------------------------------------------#

if __name__ == "__main__":
    app = QtGui.QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())
